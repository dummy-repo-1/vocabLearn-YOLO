# Importing libraries
import streamlit as st
import cv2
import numpy as np
from PIL import Image, ImageOps
import os
import torch
from ultralytics import YOLO
import easyocr
from deep_translator import GoogleTranslator
from docx import Document
import urllib.request



# import streamlit as st
# import cv2
# import numpy as np
# from PIL import Image, ImageOps
# import os
# import torch
# from ultralytics import YOLO
# import easyocr
# from deep_translator import GoogleTranslator
# from docx import Document
# import urllib.request

MODEL_URL = "https://github.com/dummy-repo-1/vocabLearn-YOLO/releases/download/weights.v1.0/bestFinal.pt"
MODEL_PATH = "bestFinal.pt"

st.set_page_config(page_title="Image Input", layout="wide")

st.title("📸 VocabLearn - A CNN based Computer Visison framework for automated word translation")
st.html(
    "<div style=font-size:20px;>Accelerate your language learning with a state-of-the-art CNN computer vision framework. Simply upload an image of your document, and our system will isolate your highlighted words, run high-accuracy OCR, translate the text, and export your new vocabulary list directly to a Word file.<br/>You can check out the source code <a href='https://github.com/projjal2025/vocabLearn', target='_blank'>Click here</a></div>"
)
st.header("Select how you want to provide the document image.")

# --- FEATURE 1: Language Selection ---
LANGUAGES = {
    'Hindi': 'hi',
    'Bengali': 'bn',
    'Spanish': 'es',
    'French': 'fr',
    'German': 'de',
    'Japanese': 'ja'
}
selected_lang_name = st.selectbox("Select Target Language for Translation:", list(LANGUAGES.keys()))
selected_lang_code = LANGUAGES[selected_lang_name]

input_method = st.radio("Choose input method:", ("Upload Image", "Capture from Camera"))

raw_image_file = []
camera_input = None

# --- METHOD 1: UPLOAD ---
if input_method == "Upload Image":
    uploaded_files = st.file_uploader("Choose an image...", 
                                        type=["png", "jpg", "jpeg"],
                                        accept_multiple_files=True)
    if uploaded_files:
        raw_image_file.extend(uploaded_files)

# --- METHOD 2: CAMERA ---
elif input_method == "Capture from Camera":
    camera_input = st.camera_input("Take a picture of the document")
    if camera_input:
        st.image(camera_input, caption="Ready for processing", use_container_width=True)
        raw_image_file.append(camera_input)

# --- LOADING MODEL ---
@st.cache_resource
def load_model():
    # Check if the file exists locally
    # if not os.path.exists(MODEL_PATH):
    # st.info("Downloading model weights for the first time. This may take a minute...")
    urllib.request.urlretrieve(MODEL_URL, MODEL_PATH)
    # st.success("Model downloaded successfully!")
        
    return YOLO(MODEL_PATH)

model = load_model()

# --- PIPELINE FUNCTIONS ---
def predict_boxes(raw_images):
    real_img = "real.jpg"
    sharpen_kernel = np.array([
        [0, -1, 0],
        [-1, 5, -1],
        [0, -1, 0]
    ])
    reader = easyocr.Reader(['en'], gpu=False)
    ocr_results = {}

    for img_file in raw_images:
        image = Image.open(img_file).convert('RGB')
        image = ImageOps.exif_transpose(image)
        image.save(real_img)
        img = cv2.imread(real_img)
        img_rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)

        results = model.predict(real_img, save=False, show_labels=False, conf=0.5)
        boxes = results[0].boxes.xyxy
        text_list = []

        for i, box in enumerate(boxes):
            x_min, y_min, x_max, y_max = map(int, box.cpu().numpy())
            h, w = img_rgb.shape[:2]
            x_min = max(0, min(x_min, w-1))
            x_max = max(0, min(x_max, w-1))
            y_min = max(0, min(y_min, h-1))
            y_max = max(0, min(y_max, h-1))

            crop = img_rgb[y_min : y_max, x_min : x_max]
            
            if crop.size == 0:
                continue

            sharp_crop = cv2.filter2D(crop, -1, sharpen_kernel)
            final_crop = cv2.bilateralFilter(sharp_crop, d=9, sigmaColor=75, sigmaSpace=75)

            ocr_output = reader.readtext(final_crop)
            for det in ocr_output:
                # Ensure the extracted text is a string
                extracted_text = str(det[1]).strip()
                if extracted_text: # Only append if it's not empty
                    text_list.append(extracted_text)

        ocr_results[img_file.name] = text_list

    return ocr_results

def translate(ocr_results, target_lang):
    translation = {}
    GTrans = GoogleTranslator(source='auto', target=target_lang)
    for key in ocr_results.keys():
        translation[key] = {}
        for item in ocr_results[key]:
            try:
                meaning = GTrans.translate(item)
                # Fallback to original text if translation returns None
                translation[key][item] = meaning if meaning else item 
            except Exception as e:
                # Handle potential translation API errors gracefully
                translation[key][item] = item 
    return translation

def write_doc(translation, target_name, new_doc=True):
    doc_filename = 'translations.docx'
    if new_doc:
        doc = Document()
        doc.add_heading(f"English-{target_name} Word Translation", level=1)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "English"
        hdr_cells[1].text = target_name

        for key in translation.keys():
            for item in translation[key].keys():
                row_cells = table.add_row().cells
                # THE FIX: Wrap assignments in str() to prevent python-docx crashes
                row_cells[0].text = str(item)
                row_cells[1].text = str(translation[key][item])

        doc.save(doc_filename)
        return doc_filename


if len(raw_image_file) > 0:
    if st.button("Process Document(s)", type="primary"):
        with st.spinner("Extracting text and translating... This might take a moment."):
            
            try:
                ocr_results = predict_boxes(raw_image_file)
                translation = translate(ocr_results, selected_lang_code)
                doc_file_path = write_doc(translation, selected_lang_name, new_doc=True)
                
                st.success("Processing complete!")

                # --- FEATURE 2: Download Button ---
                with open(doc_file_path, "rb") as file:
                    st.download_button(
                        label="⬇️ Download Translated Word Document",
                        data=file,
                        file_name="Document_Translations.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            except Exception as e:
                st.error(f"An error occurred during processing: {e}")
